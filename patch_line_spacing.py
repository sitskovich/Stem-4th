"""
Patch lessons 43-51 guided notes:
  - 3 blank lines after each guided note question
  - 5 blank lines under Key Facts
  - 5 blank lines under Q&A
Works by scanning each file dynamically, counting existing blanks,
and inserting the deficit.
"""
from docx import Document
from docx.oxml.ns import qn
import copy

LESSONS = [43, 44, 46, 47, 48, 50, 51]
BASE = r'c:\Users\TestRun\OneDrive\Desktop\STEM 4th\Lesson {n} - Guided Notes.docx'

def get_text(tbl, idx):
    tc = tbl.rows[idx]._tr.findall(qn('w:tc'))[0]
    return ''.join(t.text for t in tc.findall('.//' + qn('w:t')) if t.text).strip()

def clone_blank(tbl, ref_idx):
    """Deep-copy a blank row from ref_idx, wipe text & list formatting."""
    new_tr = copy.deepcopy(tbl.rows[ref_idx]._tr)
    for t in new_tr.findall('.//' + qn('w:t')):
        t.text = ''
    for pPr in new_tr.findall('.//' + qn('w:pPr')):
        for tag in [qn('w:numPr'), qn('w:ind')]:
            el = pPr.find(tag)
            if el is not None:
                pPr.remove(el)
    return new_tr

def insert_blanks_after(tbl, row_idx, count, blank_source_idx):
    """Insert `count` blank rows immediately after row_idx."""
    ref_tr = tbl.rows[row_idx]._tr
    for _ in range(count):
        new_tr = clone_blank(tbl, blank_source_idx)
        ref_tr.addnext(new_tr)
        ref_tr = new_tr  # chain: each new row goes after the last inserted

def classify(text):
    t = text.strip()
    if not t:
        return 'blank'
    if t.startswith('Key Facts'):
        return 'key_facts'
    if t.startswith('Q & A') or t.startswith('Q&A'):
        return 'qa'
    # topic rows start with "1.  " .. "5.  "
    if len(t) >= 2 and t[0].isdigit() and t[1] == '.':
        return 'topic'
    return 'other'

for n in LESSONS:
    path = BASE.format(n=n)
    doc = Document(path)
    tbl = doc.tables[0]

    # Find a known blank row to use as clone source
    blank_src = next(
        (i for i in range(len(tbl.rows)) if not get_text(tbl, i)),
        None
    )

    # Build row classifications (re-scan after each insertion)
    # We'll do it in a single pass collecting (idx, type), then process in reverse
    rows_info = [(i, classify(get_text(tbl, i))) for i in range(len(tbl.rows))]

    # Collect action list: (anchor_row_idx, current_blanks_after, target_blanks)
    actions = []
    i = 0
    while i < len(rows_info):
        idx, rtype = rows_info[i]
        if rtype in ('topic', 'key_facts', 'qa'):
            target = 3 if rtype == 'topic' else 5
            # Count consecutive blanks that follow
            blanks = 0
            j = i + 1
            while j < len(rows_info) and rows_info[j][1] == 'blank':
                blanks += 1
                j += 1
            deficit = target - blanks
            if deficit > 0:
                actions.append((idx, deficit))
        i += 1

    # Process in REVERSE order so insertion doesn't shift earlier indices
    for anchor_idx, deficit in reversed(actions):
        insert_blanks_after(tbl, anchor_idx, deficit, blank_src)

    doc.save(path)
    summary = [(classify(get_text(doc.tables[0], i)), i) for i in range(len(doc.tables[0].rows))]
    print(f'Lesson {n}: {len(actions)} sections patched -> saved.')

print('\nAll done!')
